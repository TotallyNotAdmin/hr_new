from docx import Document
from fastapi import APIRouter, Request, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from schemas import CreateRequest
from auth import require_role
from models import Role, RequestStatus
import io
from openpyxl import Workbook
from openpyxl.styles import Border, Side, Alignment

router = APIRouter(prefix="/requests")
STATUS_RU = {"DRAFT": "Черновик", "ON_APPROVAL": "На согласовании", "APPROVED": "Согласовано", "IN_PROGRESS": "В работе", "DONE": "Исполнено", "RETURNED": "На доработке", "REJECTED": "Отклонено"}
ACTION_RU = {"create": "Создана", "submit": "Отправлено на согласование", "approve": "Согласовано", "return": "Отправлено на доработку", "reject": "Отклонено", "take": "Взято в работу", "complete": "Завершено"}
FIELD_RU = {"employee_id": "Сотрудник", "position_id": "Штатная единица", "position_name": "Должность", "salary": "Оклад (руб.)", "address": "Адрес", "headcount": "Количество ставок", "bonus_month": "Премия (месяц, %)", "bonus_quarter": "Премия (квартал, %)", "bonus_year": "Премия (год, %)", "date": "Дата", "manager": "Админ. руководитель"}


@router.post("", summary="Создать заявку")
@router.post("/", summary="Создать заявку")
async def create_request(
        data: CreateRequest,
        request: Request,
        user=Depends(require_role([Role.MANAGER]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        req = await conn.fetchrow("""
            INSERT INTO app.requests (creator_id, scenario_type, status, comment, description)
            VALUES ($1, $2, $3, $4, $5)
            RETURNING id
        """, user["user_id"], data.scenario_type, RequestStatus.DRAFT, data.comment, data.description)

        request_id = req["id"]
        await conn.execute("""
            INSERT INTO app.request_history (request_id, action, user_id, comment)
            VALUES ($1, 'create', $2, $3)
        """, request_id, user["user_id"], data.comment)

        for field in data.fields:
            if field.name == 'manager':
                continue
            await conn.execute("""
                INSERT INTO app.request_data (request_id, field_name, as_is_value, to_be_value)
                VALUES ($1, $2, $3, $4)
            """, request_id, field.name, field.as_is, field.to_be)

        return {"request_id": request_id}


@router.post("/{request_id}/approve")
async def approve_request(
        request_id: int,
        request: Request,
        user=Depends(require_role([Role.APPROVER]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        status = await conn.fetchval(
            "SELECT status FROM app.requests WHERE id=$1",
            request_id
        )
        if status != RequestStatus.ON_APPROVAL:
            raise HTTPException(400, "Нельзя согласовать")
        await conn.execute("""
        UPDATE app.requests SET status=$1 WHERE id=$2
        """, RequestStatus.APPROVED, request_id)

        await conn.execute("""
        INSERT INTO app.request_history (request_id, action, user_id)
        VALUES ($1, 'approve', $2)
        """, request_id, user["user_id"])

        return {"status": "APPROVED"}


@router.get("/positions/{id}")
async def get_position(request: Request, id: int):
    pool = request.app.state.pool

    async with pool.acquire() as conn:
        position = await conn.fetchrow("""
            SELECT * FROM hr.positions WHERE id=$1
        """, id)

        return dict(position)


@router.get("/positions/{id}/check")
async def check_position(request: Request, id: int):
    pool = request.app.state.pool

    async with pool.acquire() as conn:
        employee = await conn.fetchrow("""
            SELECT * FROM hr.employees WHERE position_id=$1
        """, id)

        return {
            "has_employee": bool(employee)
        }


@router.post("/{request_id}/submit")
async def submit_request(
        request_id: int,
        request: Request,
        user=Depends(require_role([Role.MANAGER]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        status = await conn.fetchval("SELECT status FROM app.requests WHERE id=$1", request_id)

        if status not in (RequestStatus.DRAFT, RequestStatus.RETURNED):
            raise HTTPException(400, "Можно отправить только из DRAFT или RETURNED!")

        await conn.execute("""
            UPDATE app.requests SET status=$1, updated_at=NOW() WHERE id=$2
        """, RequestStatus.ON_APPROVAL, request_id)

        await conn.execute("""
        INSERT INTO app.request_history (request_id, action, user_id)
        VALUES ($1, 'submit', $2)
        """, request_id, user["user_id"])

        return {"status": "ON_APPROVAL"}


@router.post("/{request_id}/return")
async def return_request(
    request_id: int,
    comment: str,
    request: Request,
    user=Depends(require_role([Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool

    async with pool.acquire() as conn:
        await conn.execute("""
            UPDATE app.requests SET status=$1 WHERE id=$2
        """, RequestStatus.RETURNED, request_id)

        await conn.execute("""
            INSERT INTO app.request_history (request_id, action, user_id, comment)
            VALUES ($1, 'return', $2, $3)
        """, request_id, user["user_id"], comment)

    return {"status": "RETURNED"}


@router.post("/{request_id}/reject")
async def reject_request(
    request_id: int,
    comment: str,
    request: Request,
    user=Depends(require_role([Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool

    async with pool.acquire() as conn:
        await conn.execute("""
            UPDATE app.requests SET status=$1 WHERE id=$2
        """, RequestStatus.REJECTED, request_id)

        await conn.execute("""
            INSERT INTO app.request_history (request_id, action, user_id, comment)
            VALUES ($1, 'reject', $2, $3)
        """, request_id, user["user_id"], comment)

    return {"status": "REJECTED"}


@router.post("/{request_id}/take")
async def take_request(
        request_id: int,
        request: Request,
        user=Depends(require_role([Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        status = await conn.fetchval(
            "SELECT status FROM app.requests WHERE id=$1",
            request_id
        )
        if status != RequestStatus.APPROVED:
            raise HTTPException(400, "Можно взять только APPROVED")
        await conn.execute("""
        UPDATE app.requests SET status=$1 WHERE id=$2
        """, RequestStatus.IN_PROGRESS, request_id)

        await conn.execute("""
        INSERT INTO app.request_history (request_id, action, user_id)
        VALUES ($1, 'take', $2)
        """, request_id, user["user_id"])

        return {"status": "IN_PROGRESS"}


@router.post("/{request_id}/complete")
async def complete_request(
        request_id: int,
        request: Request,
        user=Depends(require_role([Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        status = await conn.fetchval(
            "SELECT status FROM app.requests WHERE id=$1",
            request_id
        )
        if status != RequestStatus.IN_PROGRESS:
            raise HTTPException(400, "Можно завершить только IN_PROGRESS")
        await conn.execute("""
        UPDATE app.requests SET status=$1 WHERE id=$2
        """, RequestStatus.DONE, request_id)

        await conn.execute("""
        INSERT INTO app.request_history (request_id, action, user_id)
        VALUES ($1, 'complete', $2)
        """, request_id, user["user_id"])

        return {"status": "DONE"}


@router.get("", summary="Получить список заявок")
@router.get("/", summary="Получить список заявок")
async def get_requests(
        request: Request,
        user=Depends(require_role([Role.MANAGER, Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        base_select = """
            SELECT r.id, r.creator_id, r.scenario_type, r.status, r.comment as manager_comment,
                   r.created_at, r.updated_at, r.description,
                   rh.comment as history_comment, rh.action as last_action
            FROM app.requests r
            LEFT JOIN LATERAL (
                SELECT comment, action, created_at
                FROM app.request_history
                WHERE request_id = r.id AND action IN ('return', 'reject')
                ORDER BY created_at DESC LIMIT 1
            ) rh ON true
        """

        if user["role"] == Role.MANAGER:
            # Менеджер видит только свои заявки
            query = base_select + " WHERE r.creator_id = $1 ORDER BY r.created_at DESC"
            rows = await conn.fetch(query, user["user_id"])

        elif user["role"] == Role.APPROVER:
            # Согласующий видит: ожидающие согласования + с чем он взаимодействовал + исполненные, которые согласовал
            query = base_select + """
                WHERE r.status = 'ON_APPROVAL'
                OR r.id IN (SELECT request_id FROM app.request_history WHERE user_id = $1 AND action IN ('approve','return','reject'))
                ORDER BY r.created_at DESC
            """
            rows = await conn.fetch(query, user["user_id"])

        else:  # EXECUTOR
            # Исполнитель видит: пул согласованных (ещё не взятых) + только те, которые взял/обработал он сам
            # Взятая заявка полностью скрывается от других исполнителей
            query = base_select + """
                WHERE (r.status = 'APPROVED' AND r.id NOT IN (SELECT request_id FROM app.request_history WHERE action = 'take'))
                OR r.id IN (SELECT request_id FROM app.request_history WHERE user_id = $1 AND action IN ('take','complete','return','reject'))
                ORDER BY r.created_at DESC
            """
            rows = await conn.fetch(query, user["user_id"])

        return [dict(row) for row in rows]


@router.get("/{request_id}/export")
async def export_request(
        request_id: int,
        request: Request,
        format: str = Query("xlsx", pattern="^(xlsx|docx)$"),
        user=Depends(require_role([Role.MANAGER, Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        status = await conn.fetchval("SELECT status FROM app.requests WHERE id=$1", request_id)
        if not status: raise HTTPException(404, "Заявка не найдена")
        if status == "DRAFT": raise HTTPException(400, "Экспорт черновиков запрещён")

        req = await conn.fetchrow(
            "SELECT r.*, u.full_name as creator_name FROM app.requests r JOIN app.users u ON r.creator_id = u.id WHERE r.id=$1",
            request_id)
        fields = await conn.fetch("SELECT * FROM app.request_data WHERE request_id=$1", request_id)
        history = await conn.fetch(
            "SELECT rh.action, rh.comment, rh.created_at, u.full_name as user_name FROM app.request_history rh JOIN app.users u ON rh.user_id = u.id WHERE rh.request_id=$1 ORDER BY rh.created_at",
            request_id)

        if format == "xlsx": return await _export_xlsx(req, fields, history)
        return await _export_docx(req, fields, history)


@router.get("/{request_id}")
async def get_request_detail(
    request_id: int,
    request: Request,
    user=Depends(require_role([Role.MANAGER, Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        row = await conn.fetchrow("""
        SELECT r.id, r.creator_id, r.scenario_type, r.status, r.comment as manager_comment,
        r.created_at, r.updated_at, r.description,
        rh.comment as history_comment, rh.action as last_action
        FROM app.requests r
        LEFT JOIN LATERAL (
            SELECT comment, action, created_at
            FROM app.request_history
            WHERE request_id = r.id
            AND action IN ('return', 'reject')
            ORDER BY created_at DESC
            LIMIT 1
        ) rh ON true
        WHERE r.id=$1
        """, request_id)
        if not row:
            raise HTTPException(404, "Заявка не найдена")
        if user["role"] == Role.MANAGER and row["creator_id"] != user["user_id"]:
            raise HTTPException(403, "Нет доступа")
        return dict(row)


@router.get("/{request_id}/fields")
async def get_request_fields(
        request_id: int,
        request: Request,
        user=Depends(require_role([Role.MANAGER, Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        rows = await conn.fetch("""
        SELECT * FROM app.request_data WHERE request_id=$1
        """, request_id)
        return [dict(row) for row in rows]


@router.get("/{request_id}/history")
async def get_request_history(
    request_id: int,
    request: Request,
    user=Depends(require_role([Role.MANAGER, Role.APPROVER, Role.EXECUTOR]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        rows = await conn.fetch("""
        SELECT rh.*, u.full_name, u.role
        FROM app.request_history rh
        LEFT JOIN app.users u ON rh.user_id = u.id
        WHERE rh.request_id=$1
        ORDER BY rh.created_at DESC
        """, request_id)
        return [dict(row) for row in rows]


@router.put("/{request_id}")
async def update_request(
        request_id: int,
        data: CreateRequest,
        request: Request,
        user=Depends(require_role([Role.MANAGER]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        status = await conn.fetchval("SELECT status FROM app.requests WHERE id=$1", request_id)

        if status not in (RequestStatus.DRAFT, RequestStatus.RETURNED):
            raise HTTPException(400, detail="Можно редактировать только заявки в статусе DRAFT или RETURNED")

        creator = await conn.fetchval("SELECT creator_id FROM app.requests WHERE id=$1", request_id)
        if creator != user["user_id"]:
            raise HTTPException(403, detail="Нет доступа к редактированию этой заявки")

        # Сбрасываем статус в DRAFT, если заявка была возвращена на доработку
        new_status = RequestStatus.DRAFT if status == RequestStatus.RETURNED else status

        await conn.execute("""
            UPDATE app.requests SET comment=$1, description=$2, status=$3 WHERE id=$4
        """, data.comment, data.description, new_status, request_id)

        # Обновляем заголовок в истории создания
        await conn.execute("""
            UPDATE app.request_history SET comment=$1 WHERE request_id=$2 AND action='create'
        """, data.comment, request_id)

        await conn.execute("DELETE FROM app.request_data WHERE request_id=$1", request_id)
        for field in data.fields:
            await conn.execute("""
                INSERT INTO app.request_data (request_id, field_name, as_is_value, to_be_value)
                VALUES ($1, $2, $3, $4)
            """, request_id, field.name, field.as_is, field.to_be)

        return {"status": "updated", "request_id": request_id, "new_status": new_status}


@router.delete("/{request_id}")
async def delete_request(
    request_id: int,
    request: Request,
    user=Depends(require_role([Role.MANAGER]))
):
    pool = request.app.state.pool
    async with pool.acquire() as conn:
        row = await conn.fetchrow(
            "SELECT status FROM app.requests WHERE id=$1 AND creator_id=$2",
            request_id, user["user_id"]
        )
        if not row:
            raise HTTPException(404, "Заявка не найдена или нет прав")
        if row["status"] != RequestStatus.DRAFT:
            raise HTTPException(400, "Можно удалить только черновики")

        # Проверка: отправлялась ли заявка ранее хоть раз
        was_submitted = await conn.fetchval(
            "SELECT EXISTS(SELECT 1 FROM app.request_history WHERE request_id=$1 AND action='submit')",
            request_id
        )
        if was_submitted:
            raise HTTPException(403, "Невозможно удалить заявку, которая уже отправлялась")

        # Удаление (CASCADE)
        await conn.execute("DELETE FROM app.requests WHERE id=$1", request_id)
        return {"message": "Заявка удалена"}


# --- Функции экспорта ---
def _apply_xlsx_format(ws):
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical='center')
            if cell.value is not None: max_length = max(max_length, len(str(cell.value)))
        if col_letter: ws.column_dimensions[col_letter].width = min(max(max_length + 2, 10), 50)


async def _export_xlsx(req, fields, history):
    wb = Workbook()
    ws = wb.active
    ws.title = f"Заявка #{req['id']}"
    ws.append(["ID", req["id"], "Создатель", req["creator_name"]])
    ws.append(["Сценарий", req["scenario_type"], "Статус", STATUS_RU.get(req["status"], req["status"])])
    ws.append(["Дата создания", req["created_at"].strftime("%d.%m.%Y %H:%M")])
    ws.append(["Комментарий", req.get("manager_comment") or "—"])
    ws.append(["Описание", req.get("description") or "—"])
    ws.append([])
    ws.append(["Поле", "AS IS", "TO BE"])
    for f in fields: ws.append([FIELD_RU.get(f["field_name"], f["field_name"]), f["as_is_value"] or "—", f["to_be_value"] or "—"])
    ws.append([])
    ws.append(["Дата", "Действие", "Исполнитель", "Комментарий"])
    for h in history: ws.append([h["created_at"].strftime("%d.%m.%Y %H:%M"), ACTION_RU.get(h["action"], h["action"]), h["user_name"], h.get("comment") or "—"])
    _apply_xlsx_format(ws)
    output = io.BytesIO(); wb.save(output); output.seek(0)
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment; filename=request_{req['id']}.xlsx"})


async def _export_docx(req, fields, history):
    doc = Document()
    doc.add_heading(f"Заявка #{req['id']}", 0)
    doc.add_paragraph(f"Статус: {STATUS_RU.get(req['status'], req['status'])} | Создатель: {req['creator_name']}")
    doc.add_paragraph(f"Сценарий: {req['scenario_type']} | Дата: {req['created_at'].strftime('%d.%m.%Y %H:%M')}")
    if req.get("manager_comment"): doc.add_paragraph(f"Комментарий: {req['manager_comment']}")
    if req.get("description"): doc.add_paragraph(f"Описание: {req['description']}")
    doc.add_heading("Данные заявки", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
    for i, t in enumerate(["Поле", "AS IS", "TO BE"]): table.rows[0].cells[i].text = t
    for f in fields:
        row = table.add_row().cells
        row[0].text = FIELD_RU.get(f["field_name"], f["field_name"])
        row[1].text = f["as_is_value"] or "—"
        row[2].text = f["to_be_value"] or "—"
    doc.add_heading("История действий", level=1)
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
    for i, t in enumerate(["Дата", "Действие", "Исполнитель", "Комментарий"]): table.rows[0].cells[i].text = t
    for h in history:
        row = table.add_row().cells
        row[0].text = h["created_at"].strftime("%d.%m.%Y %H:%M")
        row[1].text = ACTION_RU.get(h["action"], h["action"])
        row[2].text = h["user_name"]
        row[3].text = h.get("comment") or "—"
    output = io.BytesIO(); doc.save(output); output.seek(0)
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=request_{req['id']}.docx"})
